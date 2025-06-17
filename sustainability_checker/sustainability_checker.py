# import streamlit as st
# import pickle
# import docx
# import PyPDF2
# import re
# import numpy as np
# import pandas as pd
# import nltk
# from nltk.corpus import stopwords
# from nltk.stem import WordNetLemmatizer

# st.set_page_config(page_title="Resume Sustainability Score", page_icon="🌱", layout="centered")

# # === Load NLTK Resources ===
# @st.cache_resource
# def load_nltk_resources():
#     nltk.download('stopwords')
#     nltk.download('wordnet')
#     nltk.download('omw-1.4')
#     return set(stopwords.words('english')), WordNetLemmatizer()

# stop_words, lemmatizer = load_nltk_resources()

# # === Load Model Files ===
# @st.cache_resource

# def load_model_files():
#     try:
#         model = pickle.load(open('sustainability_model.pkl', 'rb'))
#         tfidf = pickle.load(open('tfidf_vectorizer.pkl', 'rb'))
#         ohe = pickle.load(open('one_hot_encoder.pkl', 'rb'))
#         return model, tfidf, ohe
#     except Exception as e:
#         st.error(f"Error loading model files: {e}")
#         st.stop()

# model, tfidf, ohe = load_model_files()

# # === Text Extraction ===
# def extract_text_from_pdf(file):
#     reader = PyPDF2.PdfReader(file)
#     return ''.join(page.extract_text() for page in reader.pages if page.extract_text())

# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return '\n'.join(p.text for p in doc.paragraphs)

# def extract_text_from_txt(file):
#     try:
#         return file.read().decode('utf-8')
#     except UnicodeDecodeError:
#         return file.read().decode('latin-1')

# def handle_uploaded_file(file):
#     ext = file.name.split('.')[-1].lower()
#     if ext == 'pdf':
#         return extract_text_from_pdf(file)
#     elif ext == 'docx':
#         return extract_text_from_docx(file)
#     elif ext == 'txt':
#         return extract_text_from_txt(file)
#     else:
#         raise ValueError("Unsupported file type.")

# # === Resume Cleaning ===
# def clean_resume(text):
#     text = re.sub(r"http\S+", ' ', text)
#     text = re.sub(r"RT|cc", ' ', text)
#     text = re.sub(r"#\S+", ' ', text)
#     text = re.sub(r"@\S+", ' ', text)
#     text = text.lower()
#     text = re.sub(r"[%s]" % re.escape(r"""!#$%&'()*+,-./:;<=>?@[\\]^_`{|}~"""), ' ', text)
#     text = re.sub(r"[^\x00-\x7f]", ' ', text)
#     text = re.sub(r"\s+", ' ', text).strip()
#     words = text.split()
#     text = ' '.join(lemmatizer.lemmatize(w) for w in words if w not in stop_words)
#     return text

# # === Prediction ===
# def predict_sustainability_score(text, category):
#     cleaned = clean_resume(text)
#     resume_vec = tfidf.transform([cleaned]).toarray()
#     cat_vec = ohe.transform(pd.DataFrame([[category]], columns=['Category'])).toarray()
#     final_input = np.hstack((cat_vec, resume_vec))
#     return model.predict(final_input)[0]

# # === Main UI ===




# def run_sustainability_checker():
#     st.title("🌿 Resume Sustainability Score Predictor")
#     st.markdown("Upload resumes (PDF, DOCX, or TXT) and select job categories to get scores.")

#     uploaded_files = st.file_uploader("Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)

#     sample_categories = ohe.categories_[0].tolist()
#     selected_category = st.selectbox("Select Job Category", sample_categories)

#     results = []
#     MAX_FILE_SIZE_MB = 5

#     if uploaded_files:
#         for i, file in enumerate(uploaded_files):
#             st.divider()
#             st.subheader(f"📄 {file.name}")

#             if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
#                 st.warning(f"⚠️ {file.name} exceeds {MAX_FILE_SIZE_MB}MB and was skipped.")
#                 continue

#             try:
#                 resume_text = handle_uploaded_file(file)
#                 st.success("Resume text extracted successfully.")

#                 if st.checkbox(f"Show Extracted Text - {file.name}", value=False):
#                     preview = resume_text[:3000] + ('...' if len(resume_text) > 3000 else '')
#                     st.text_area("Extracted Text", preview, height=300)

#                 score = predict_sustainability_score(resume_text, selected_category)

#                 if score >= 75:
#                     feedback = "✅ Strong Resume!"
#                 elif score >= 50:
#                     feedback = "⚠️ Moderate – can be improved."
#                 else:
#                     feedback = "❌ Needs significant improvement."

#                 st.markdown(f"### 🌱 Sustainability Score: **{round(score, 2)} / 100**")
#                 st.markdown(f"**Feedback:** {feedback}")

#                 results.append({
#                     "Filename": file.name,
#                     "Selected Category": selected_category,
#                     "Sustainability Score": round(score, 2)
#                 })

#             except Exception as e:
#                 st.error(f"❌ Error processing {file.name}: {e}")

#         if results:
#             st.divider()
#             st.subheader("📊 Summary of Scores")
#             df_results = pd.DataFrame(results)
#             st.dataframe(df_results)

#             csv = df_results.to_csv(index=False).encode('utf-8')
#             st.download_button(
#                 label="📥 Download Results as CSV",
#                 data=csv,
#                 file_name="sustainability_scores.csv",
#                 mime='text/csv'
#             )




# Required installations:
# pip install streamlit scikit-learn python-docx PyPDF2

import streamlit as st
import pickle
import docx
import PyPDF2
import re
import numpy as np
import pandas as pd
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import os

st.set_page_config(page_title="Resume Sustainability Score", page_icon="🌱", layout="centered")

# === Load NLTK Resources ===
@st.cache_resource
def load_nltk_resources():
    nltk.download('stopwords')
    nltk.download('wordnet')
    nltk.download('omw-1.4')
    return set(stopwords.words('english')), WordNetLemmatizer()

stop_words, lemmatizer = load_nltk_resources()

# === Load Model Files ===
@st.cache_resource
def load_model_files():
    try:
        base_path = os.path.dirname(__file__)
        model = pickle.load(open(os.path.join(base_path, 'sustainability_model.pkl'), 'rb'))
        tfidf = pickle.load(open(os.path.join(base_path, 'tfidf_vectorizer.pkl'), 'rb'))
        ohe = pickle.load(open(os.path.join(base_path, 'one_hot_encoder.pkl'), 'rb'))
        return model, tfidf, ohe
    except Exception as e:
        st.error(f"Error loading model files: {e}")
        st.stop()

model, tfidf, ohe = load_model_files()

# === Text Extraction ===
def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    return ''.join(page.extract_text() for page in reader.pages if page.extract_text())

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return '\n'.join(p.text for p in doc.paragraphs)

def extract_text_from_txt(file):
    try:
        return file.read().decode('utf-8')
    except UnicodeDecodeError:
        return file.read().decode('latin-1')

def handle_uploaded_file(file):
    ext = file.name.split('.')[-1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(file)
    elif ext == 'docx':
        return extract_text_from_docx(file)
    elif ext == 'txt':
        return extract_text_from_txt(file)
    else:
        raise ValueError("Unsupported file type.")

# === Resume Cleaning ===
def clean_resume(text):
    text = re.sub(r"http\S+", ' ', text)
    text = re.sub(r"RT|cc", ' ', text)
    text = re.sub(r"#\S+", ' ', text)
    text = re.sub(r"@\S+", ' ', text)
    text = text.lower()
    text = re.sub(r"[%s]" % re.escape(r"""!#$%&'()*+,-./:;<=>?@[\\]^_`{|}~"""), ' ', text)
    text = re.sub(r"[^\x00-\x7f]", ' ', text)
    text = re.sub(r"\s+", ' ', text).strip()
    words = text.split()
    text = ' '.join(lemmatizer.lemmatize(w) for w in words if w not in stop_words)
    return text

# === Prediction ===
def predict_sustainability_score(text, category):
    cleaned = clean_resume(text)
    resume_vec = tfidf.transform([cleaned]).toarray()
    cat_vec = ohe.transform(pd.DataFrame([[category]], columns=['Category'])).toarray()
    final_input = np.hstack((cat_vec, resume_vec))
    return model.predict(final_input)[0]

# === Main UI ===
def run_sustainability_checker():
    st.title("🌿 Resume Sustainability Score Predictor")
    st.markdown("Upload resumes (PDF, DOCX, or TXT) and select job categories to get scores.")

    uploaded_files = st.file_uploader("Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True)

    sample_categories = ohe.categories_[0].tolist()
    selected_category = st.selectbox("Select Job Category", sample_categories)

    results = []
    MAX_FILE_SIZE_MB = 5

    if uploaded_files:
        for i, file in enumerate(uploaded_files):
            st.divider()
            st.subheader(f"📄 {file.name}")

            if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
                st.warning(f"⚠️ {file.name} exceeds {MAX_FILE_SIZE_MB}MB and was skipped.")
                continue

            try:
                resume_text = handle_uploaded_file(file)
                st.success("Resume text extracted successfully.")

                if st.checkbox(f"Show Extracted Text - {file.name}", value=False):
                    preview = resume_text[:3000] + ('...' if len(resume_text) > 3000 else '')
                    st.text_area("Extracted Text", preview, height=300)

                score = predict_sustainability_score(resume_text, selected_category)

                if score >= 75:
                    feedback = "✅ Strong Resume!"
                elif score >= 50:
                    feedback = "⚠️ Moderate – can be improved."
                else:
                    feedback = "❌ Needs significant improvement."

                st.markdown(f"### 🌱 Sustainability Score: **{round(score, 2)} / 100**")
                st.markdown(f"**Feedback:** {feedback}")

                results.append({
                    "Filename": file.name,
                    "Selected Category": selected_category,
                    "Sustainability Score": round(score, 2)
                })

            except Exception as e:
                st.error(f"❌ Error processing {file.name}: {e}")

        if results:
            st.divider()
            st.subheader("📊 Summary of Scores")
            df_results = pd.DataFrame(results)
            st.dataframe(df_results)

            csv = df_results.to_csv(index=False).encode('utf-8')
            st.download_button(
                label="📅 Download Results as CSV",
                data=csv,
                file_name="sustainability_scores.csv",
                mime='text/csv'
            )